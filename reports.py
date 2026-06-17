# reports.py
# DOCX- og PDF-rapportgenerering for byggTotal.

import io
from datetime import datetime

import pandas as pd

try:
    from docx import Document
except Exception:
    Document = None

try:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
except Exception:
    SimpleDocTemplate = None


def build_docx_report(summary_dict, material_summary, extra_sections=None):
    if Document is None:
        return None
    doc = Document()
    doc.add_heading("byggTotal – Prosjektrapport", 0)
    p = doc.add_paragraph()
    p.add_run("Generert: ").bold = True
    p.add_run(datetime.now().strftime("%d.%m.%Y %H:%M"))

    doc.add_heading("Prosjektoversikt", level=1)
    for k, v in summary_dict.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Materialoversikt", level=1)
    table = doc.add_table(rows=1, cols=len(material_summary.columns))
    table.style = "Table Grid"
    for i, col in enumerate(material_summary.columns):
        table.rows[0].cells[i].text = str(col)
    for _, row in material_summary.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = f"{val}" if not isinstance(val, float) else f"{val:,.2f}".replace(",", " ")

    if extra_sections:
        for title, df in extra_sections:
            if df is None or df.empty:
                continue
            doc.add_heading(title, level=1)
            t = doc.add_table(rows=1, cols=len(df.columns))
            t.style = "Table Grid"
            for i, col in enumerate(df.columns):
                t.rows[0].cells[i].text = str(col)
            for _, row in df.iterrows():
                cells = t.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = f"{val}" if not isinstance(val, float) else f"{val:,.2f}".replace(",", " ")

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def build_pdf_report(summary_dict, material_summary, extra_sections=None):
    if SimpleDocTemplate is None:
        return None
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    elements = [
        Paragraph("byggTotal – Prosjektrapport", styles["Title"]),
        Spacer(1, 12),
        Paragraph(f"Generert: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles["Normal"]),
        Spacer(1, 12),
    ]

    summary_table_data = [["Parameter", "Verdi"]] + [[str(k), str(v)] for k, v in summary_dict.items()]
    t1 = Table(summary_table_data, hAlign="LEFT")
    t1.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements += [Paragraph("Prosjektoversikt", styles["Heading2"]), t1, Spacer(1, 16)]

    t2 = Table([list(material_summary.columns)] + material_summary.astype(str).values.tolist(), hAlign="LEFT")
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    elements += [Paragraph("Materialoversikt", styles["Heading2"]), t2, Spacer(1, 16)]

    if extra_sections:
        for title, df in extra_sections:
            if df is None or df.empty:
                continue
            t = Table([list(df.columns)] + df.astype(str).values.tolist(), hAlign="LEFT")
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            elements += [Paragraph(title, styles["Heading2"]), t, Spacer(1, 16)]

    doc.build(elements)
    buffer.seek(0)
    return buffer.getvalue()


def make_report_summary_dict(filename, filtered_df):
    return {
        "Fil": filename,
        "Antall elementer": int(len(filtered_df)),
        "Total lengde [m]": float(pd.to_numeric(filtered_df["Lengde [m]"], errors="coerce").fillna(0).sum()),
        "Total areal [m2]": float(pd.to_numeric(filtered_df["Areal [m2]"], errors="coerce").fillna(0).sum()),
        "Total volum [m3]": float(pd.to_numeric(filtered_df["Volum [m3]"], errors="coerce").fillna(0).sum()),
        "Total vekt [kg]": float(pd.to_numeric(filtered_df["Vekt [kg]"], errors="coerce").fillna(0).sum()),
        "Total kostnad [kr]": float(pd.to_numeric(filtered_df["Kostnad [kr]"], errors="coerce").fillna(0).sum()),
        "Total CO2 [kgCO2e]": float(pd.to_numeric(filtered_df["CO2 [kgCO2e]"], errors="coerce").fillna(0).sum()),
    }
